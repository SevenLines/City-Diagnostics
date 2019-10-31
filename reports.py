import re
from collections import OrderedDict
from datetime import datetime

from PyQt5 import QtCore

from itertools import groupby

import docx
from PyQt5.QtCore import QObject
from sqlalchemy import or_, and_, func, text, literal_column, case

from db import session, Session
from helpers import Range, check_in, set_vertical_cell_direction, RangeAvg, RangeCustom, add_row
from models import Attribute, Params, ListAttrib, Road, High, Way
from shapely.geometry import Polygon, LineString
from docxtpl import DocxTemplate


def get_km(value):
    return "{}+{:03}".format(value // 1000, value % 1000)


POSITION_LEFT = -1
POSITION_RIGHT = 1
POSITION_BOTH = 0

class SmoothMixin(object):
    def __init__(self,  *args, **kwargs) -> None:
        super().__init__()
        self.smoothnes = None
        self.default_category = kwargs.get('default_category', '4')
        self.info = None

    def get_road_type_and_category(self):
        if self.info:
            return self.info
        attributes = Attribute.query_by_road(self.session, self.road.id).filter(
            Attribute.ID_Type_Attr.in_(['2320', '2306'])  # Тип покрытия, Категория
        ).join(Params, Params.attribute_id == Attribute.id).with_entities(
            Attribute.ID_Type_Attr,
            Params.value
        )

        attributes = list(attributes)

        attributes = {
            a.ID_Type_Attr: a.value for a in attributes
        }

        self.info = {
            'category': attributes.get('2306', self.default_category),
            'type': attributes.get('2320', '???покрытие'),
        }

        return self.info

    def get_smooth_data(self):
        if self.smoothnes:
           return self.smoothnes
        self.smoothnes = []
        column = literal_column("L1 / {delta} * {delta}".format(delta=self.delta)) # literal_column("L1 / 100 * 100")

        def get_query(direction):
            return self.session.query(Attribute).filter(
                Attribute.ID_Type_Attr == 2310,
                Attribute.id.in_(
                    Attribute.query_by_road(self.session, self.road.id)
                        .join(Params)
                        .filter(Params.id == '231')
                        .filter(Params.value == direction)
                        .with_entities(Attribute.id)
                )
            ).join(Params).filter(Params.id == '233')

        forward = get_query("Прямое").with_entities(
            column.label('pos'),
            literal_column("Avg(Cast(ValueParam as float))").label('forward'),
            literal_column('0').label('backward'),
        ).group_by(column)

        backward = get_query("Обратное").with_entities(
            column,
            literal_column("0"),
            literal_column("Avg(Cast(ValueParam as float))"),
        ).group_by(column)

        qs = forward.union(backward).subquery('t')

        attributes = self.session.query(
            qs.c.pos.label('pos'),
            case([(func.max(qs.c.forward) > 0, func.max(qs.c.forward))], else_=func.max(qs.c.backward)).label('forward'),
            case([(func.max(qs.c.backward) > 0, func.max(qs.c.backward)),], else_=func.max(qs.c.forward)).label('backward'),
        ).group_by(qs.c.pos)

        info = self.get_road_type_and_category()
        max_iri = {
            '1': 4,
            '1В': 4.5,
            '1B': 4.5,
            '2': 4.5,
            '3': 5.5,
            '4': 6.5,
            '5': 8.0,
        }.get(info['category'], 5.5)

        for a in attributes:
            if self.start <= a.pos <= self.end:
                self.smoothnes.append({
                    'backward': round(a.backward, 2),
                    'forward': round(a.forward, 2),
                    'delta': self.end - a.pos if a.pos + self.delta > self.end else self.delta,
                    'pos': a.pos,
                    'max_iri': max_iri,
                    'is_bad': max(a.backward, a.forward) > max_iri
                })

        return self.smoothnes

    def fill_smooth_data(self, table):
        table.rows[0].cells[0].text = self.road.Name
        table.rows[1].cells[0].text = "«Дата»: {:%d.%m.%Y}".format(datetime.now())

        smoothnes = self.get_smooth_data()
        info = self.get_road_type_and_category()

        for a in smoothnes:
            row = add_row(table, [
                str(a['backward']),
                str(a['forward']),
                "{}".format(a['delta']),
                get_km(a['pos']),
                info['category'],
                info['type'],
            ])


class BadWellsReport(object):
    def __init__(self, session) -> None:
        self.session = session

    def __call__(self, ID_Road, *args, **kwargs):
        """
        ДЕФЕКТ: Параметр в Свойствах "Принадлежность" (ID_Param: 3939 or 3937)
        """
        DEFECTS = {
            "0602": {  # Люки смотровых колодцов
                1: "Разрушение а/б около канализационного люка",
                2: "Продавленный канализационный люк",
                3: "Отсутствующий канализационный люк",
                4: "Выпирающий канализационный люк",
            },
            "0601": {  # Решетки ливневых колодцев
                1: "Разрушение а/б возле люка ливневой канализации",
                2: "Продавлен люк ливневой канализации",
                3: "Отсутствует люк ливневой канализации",
                4: "Выпирающий люк ливневой канализации",
            }
        }

        query = Attribute.query_by_road(self.session, ID_Road) \
            .join(Params).with_entities() \
            .filter(Attribute.ID_Type_Attr.in_(["0602", "0602"])) \
            .filter(Params.id.in_(["3939", "3937"])) \
            .order_by(Attribute.L1, Attribute.L2) \
            .with_entities(
            Attribute.L1,
            Attribute.ID_Type_Attr,
            Attribute.Image_Points,
            Attribute.Image_Counts,
            Params.value.label("defect")
        )

        out = []
        for r in query:
            type = {
                "0602": "Люк смотрового колодца",
                "0601": "Решетка ливневого колодца",
            }.get(r.ID_Type_Attr)

            position = r.L1

            defects = ", ".join([DEFECTS.get(r.ID_Type_Attr)[int(d.strip())] for d in re.split("[.,]", r.defect)])

            out.append((type, position, defects, r.Image_Points, r.Image_Counts))

        return out


class BarringReport(object):
    def __init__(self, session) -> None:
        self.session = session

    def __call__(self, ID_Road, _start, _end, *args, **kwargs):
        """
        010109 -- бортовой камень, ID_Param: 3935 == 1ГП то condition плохое
        020301	Криволинейный брус (Металическое)
        020302	Сигнальный столбик
        020303	Типа Нью-Джерси
        020304	Пешеходное ограждение
        020305	Парапеты
        020306	Тросовое
        020307	Не стандартное

        :param ID_Road:
        :return:
        """
        attributes = Attribute.query_by_road(self.session, ID_Road).filter(
            Attribute.ID_Type_Attr.in_([
                "010109",  # бортовой камен
                "020301",
                "020302",
                "020303",
                "020304",
                "020305",
                "020306",
                "020307",
            ])
        ).join(ListAttrib).outerjoin(Params) \
            .filter(
            or_(
                Params.id == None,
                Params.id.in_([
                    3935,  # Марка (Тип)
                    830,  # высота
                    831,  # высота
                    833,  # высота
                    835,  # высота
                    837,  # высота
                    838,  # высота
                    2796,  # высота
                ])
            )
        ).with_entities(
            ListAttrib.name_attribute,
            Attribute.L1,
            Attribute.L2,
            Attribute.ID_Type_Attr,
            Attribute.Image_Points,
            Attribute.Image_Counts,
            Params.value
        ).order_by(Attribute.L1, Attribute.L2)

        out = []
        for r in attributes:
            start = max(_start, min(_end, max(0, r.L1)))
            end = max(_start, min(_end, r.L2))
            length = end - start

            if length == 0:
                continue

            range_ = "{} — {}".format(get_km(start), get_km(end))
            type_ = r.name_attribute
            condition = "плохое" if r.value else "хорошее"

            points = Attribute.get_points(r.Image_Points, r.Image_Counts)
            position = "Слева" if points[0].a < 0 else 'Справа'

            if length:
                out.append({
                    "Участок": range_,
                    "Позиция": position,
                    "Тип": type_,
                    "condition": condition,
                    "length": length,
                    'start': start,
                    'end': r.L2,
                    'points': [(p.x, p.y) for p in points],
                    'is_barrier': r.ID_Type_Attr != '010109'
                })

        return out


class TrailReport(object):
    def __init__(self, session) -> None:
        self.session = session

    def _get_info(self, key, items):
        info = {i.param_id: i for i in items}
        info['title'] = key
        info['L1'] = items[0].L1
        info['L2'] = items[0].L2
        points = Attribute.get_points(items[0].Image_Points, items[0].Image_Counts)
        info['x'] = points[0].x
        info['y'] = points[0].y
        return info

    def __call__(self, ID_Road, *args, **kwargs):
        attributes = Attribute.query_by_road(self.session, ID_Road).filter(
            Attribute.ID_Type_Attr == '0407'
        ).join(Params).order_by(Attribute.L1, Attribute.L2, Attribute.id).with_entities(
            Attribute.L1,
            Attribute.L2,
            Attribute.id,
            Attribute.Image_Counts,
            Attribute.Image_Points,
            Params.value,
            Params.id.label('param_id')
        )

        data = [self._get_info(key, list(items)) for key, items in groupby(attributes, key=lambda x: x.id)]
        data = [i for i in data if i.get(3204, '')]

        out = []
        for i in data:
            if i.get(3204):
                out.append({
                    "x": i['x'],
                    "y": i['y'],
                    "start": i['L1'],
                    "end": i['L2'],
                    "type": i.get(3205).value if 3205 in i else '',
                    "defect": ", ".join([{
                        '1': 'Разрушение а/б на пересекаемых трамвайных (ж/д) путях',
                        '2': 'Отклонение верха головки рельса трамвайных или железнодорожных путей, '
                             'расположенных в пределах проезжей части, относительно покрытия более 2,0 см.',
                        '3': 'Разрушение ц/б плит на пересекаемых трамвайных (ж/д) путях'
                    }.get(i) for i in re.split("[.,]", i.get(3204).value)])
                })

        return out


class DiagnosticsReport(SmoothMixin, QObject):
    DEFECTS = {
        "01020103": {
            "title": "Выбоины",
        },
        "01020104": {
            "title": "Сетка трещин",
        },
        "01020102": {
            "title": "Карты(Заплаты)",
        },
        "01020101": {
            "title": "Трещины",
        }
    }

    progressed = QtCore.pyqtSignal(int, int, str)

    def __init__(self, road_id, *args, **kwargs) -> None:
        super(DiagnosticsReport, self).__init__(*args, **kwargs)
        self.session = Session()
        self.defects = None
        self.barrier_data = None
        self.delta = kwargs.get('delta', 100)
        self.road = self.session.query(Road).get(road_id)
        self.info = None
        self.start, self.end = self.road.get_length(self.session)

    def close(self):
        self.session.close()

    @classmethod
    def calculate_delta_to_next(cls, data, max_value=None):
        preprocessed = []
        for idx, c in enumerate(data[:-1]):
            n = data[idx + 1]
            dl = abs(n[0] - c[1])

            if max_value and dl > max_value:
                dl = None

            preprocessed.append(
                (c[0], c[1], dl)
            )
        preprocessed.append(
            (data[-1][0], data[-1][1], None)
        )
        return preprocessed

    @classmethod
    def join_ranges(cls, data, list_values=None):
        if not data:
            return []
        list_values_ordered = sorted(list_values)

        preprocessed = []
        for idx, c in enumerate(data[:-1]):
            new_value = c[2]
            if new_value is not None:
                for v in list_values_ordered:
                    if new_value <= v:
                        new_value = v
                        break

            preprocessed.append(
                (c[0], c[1], new_value)
            )

        preprocessed.append((
            data[-1][0],
            data[-1][1],
            None
        ))

        ranges = []
        previous_value = preprocessed[0][2]
        current_range = preprocessed[0]

        for idx, (l1, l2, value) in enumerate(preprocessed):
            if previous_value is None:
                if value is None:
                    ranges.append((l1, l2, None))
                else:
                    current_range = (l1, l2, value)
            else:
                if value == previous_value:
                    if idx == len(preprocessed) - 1:
                        current_range = (current_range[0], l2, value)
                    else:
                        current_range = (current_range[0], l1, value)
                else:
                    ranges.append((current_range[0], l2, previous_value))
                    current_range = (l2, l2, value)

            previous_value = value

        return ranges

    def get_koleynost(self):
        qs = self.session.query(Attribute).filter(
            Attribute.id.in_(
                Attribute.query_by_road(self.session, self.road.id)
                    .filter(Attribute.ID_Type_Attr == "01020106")
                    .join(Params)
                    .with_entities(Attribute.id)
            )
        ).join(Params).filter(Params.id == 212) \
            .with_entities(Attribute.L1, Attribute.L2, Params.value) \
            .order_by(Attribute.L1, Attribute.L2)

        rng = RangeAvg(self.start, self.end)

        out = []
        for row in qs:
            rng.add_subrange(row.L1, row.L2 if row.L2 > row.L1 else row.L1 + 1, min(70, int(float(row.value) * 10)))
            out.append((row.L1, row.L2, min(70, int(float(row.value) * 10))))

        out = self.join_ranges([r for r in rng.ranges if r[2] is not None], [10, 20, 30, 40, 50, 70])
        return out

    def get_defects_report(self):

        attributes = self.session.query(
            Attribute.L1,
            Attribute.L2,
            Attribute.ID_Type_Attr,
            ListAttrib.name_attribute,
            Attribute.Image_Counts,
            Attribute.Image_Points
        )\
            .join(High)\
            .join(Way)\
            .filter(Way.road_id == self.road.id) \
            .filter(Attribute.ID_Type_Attr.in_(self.DEFECTS.keys())) \
            .outerjoin(ListAttrib, ListAttrib.id == Attribute.ID_Type_Attr) \
            .order_by(Attribute.L1, Attribute.L2)

        defects = {}
        LMax = 0
        attributes = list(attributes)
        for r in attributes:
            name = r.name_attribute
            value = 0
            position = POSITION_BOTH

            if r.ID_Type_Attr == '01020101':
                points = Attribute.get_points(r.Image_Points, r.Image_Counts)

                max_p = max([p.a for p in points])
                min_p = min([p.a for p in points])
                da = max_p - min_p

                dy = r.L2 - r.L1
                if da > dy:
                    name = "Поперечные трещины"
                else:
                    name = "Продольные трещины"
                value = LineString([(p.x, p.y) for p in points]).length
                if max_p * min_p < 0:
                    position = POSITION_BOTH
                elif max_p > 0:
                    position = POSITION_RIGHT
                else:
                    position = POSITION_LEFT
            elif r.ID_Type_Attr in ('01020102', '01020103', '01020104'):  # Выбоины, сетка трещин
                points = Attribute.get_points(r.Image_Points, r.Image_Counts)

                max_p = max([p.a for p in points])
                min_p = min([p.a for p in points])
                da = max_p - min_p

                if max_p * min_p < 0:
                    position = POSITION_BOTH
                elif max_p > 0:
                    position = POSITION_RIGHT
                else:
                    position = POSITION_LEFT

                if len(points) > 2:
                    polygon = Polygon([(p.x, p.y) for p in points])
                    value = polygon.area
                else:
                    value = LineString([(p.x, p.y) for p in points]).length * 6
                name = self.DEFECTS.get(r.ID_Type_Attr, {"title": name})['title']

            item = defects.setdefault(name, [])
            item.append((r.L1, r.L2, {"value": value, "position": position}))

            LMax = max(LMax, r.L1, r.L2)

        longitudinal_cracks = defects.get("Продольные трещины", [])
        longitudinal_cracks_ranges = []
        if longitudinal_cracks:
            rng = Range(self.start, self.end)
            for r in longitudinal_cracks:
                rng.add_subrange(r[0], r[1] if r[1] > r[0] else r[0] + 1, r[2])
            longitudinal_cracks_ranges = [i for i in rng.ranges if i[2]]

        transverse_cracks = defects.get("Поперечные трещины", [])
        transverse_cracks_ranges = []
        if transverse_cracks:
            transverse_cracks_ranges = self.join_ranges(
                self.calculate_delta_to_next(transverse_cracks, 40),
                [2, 3, 4, 6, 8, 10, 20, 40]
            )

        potholes = defects.get("Выбоины", [])
        potholes_ranges = []
        if potholes:
            potholes_ranges = self.join_ranges(
                self.calculate_delta_to_next(potholes, 20),
                [4, 10, 20]
            )

        koleynost = self.get_koleynost()

        return {
            'Поперечные трещины': transverse_cracks_ranges,
            'Продольные трещины': longitudinal_cracks_ranges,
            'Трещины': defects.get("Продольные трещины", []) + defects.get("Поперечные трещины", []),
            'Выбоины': potholes_ranges,
            'Выбоины_raw': defects.get("Выбоины", []),
            'Колейность': koleynost,
            'Сетка трещин': defects.get("Сетка трещин"),
            'Карты(Заплаты)': defects.get("Карты(Заплаты)", []),
        }

    def get_range(self):
        delta = self.delta
        ranges = range(self.start, self.end + 1, delta)
        return ranges, delta

    def get_defects(self, round_score=True, skip_koleynost=False):
        if self.defects is not None:
            return self.defects

        defects = self.get_defects_report()

        ranges, delta = self.get_range()
        self.defects = []

        for offset in ranges:
            row = {
                "address": get_km(offset),
                "pos": offset,
                "length": self.end - offset if offset + delta > self.end else delta,
                "defects": [],
                "score": 0,
            }

            defect_cell_offset = 3
            for item in defects['Продольные трещины'] or []:
                if check_in(item, (offset, offset + 100)):
                    row['defects'].append({
                        'type': 'Продольные боковые трещины',
                        'short': "+",
                        "cell": defect_cell_offset,
                        'alone': False,
                        'score': 3.5,
                        'length': item[0] - item[1],
                        'description': "продольные боковые трещины",
                        'defect_code': '11',
                    })

            defect_cell_offset = 4
            for item in defects['Поперечные трещины'] or []:
                if offset <= item[0] <= offset + delta or offset <= item[1] <= offset + delta:
                    if item[2] == None:
                        cell_offset = 0
                    else:
                        cell_offset = [40, 20, 10, 8, 6, 4, 3, 2].index(item[2]) + 1

                    defect_item = {
                        None: {"description": "более 40м", "code": '1', "score": 5.0},
                        40: {"description": "20-40м", "code": '2', "score": 4.9},
                        20: {"description": "10-20м", "code": '3', "score": 4.6},
                        10: {"description": "8-10м", "code": '4', "score": 4.3},
                        8: {"description": "6-8м", "code": '5', "score": 3.9},
                        6: {"description": "4-6м", "code": '6', "score": 3.6},
                        4: {"description": "3-4м", "code": '7', "score": 3.2},
                        3: {"description": "2-3м", "code": '8', "score": 2.9},
                        2: {"description": "1-2м", "code": '9', "score": 2.6},
                    }.get(item[2])

                    row['defects'].append({
                        'type': 'Поперечные трещины',
                        'short': "+",
                        'cell': defect_cell_offset + cell_offset,
                        'alone': item[2] is None,
                        'address': (item[0] + item[1]) / 2,
                        'length': item[1] + item[0],
                        'score': defect_item['score'],
                        'description': "поперечные трещины, на расстоянии {}".format(
                            defect_item['description'],
                        ),
                        'defect_code': defect_item['code'],
                    })

            defect_cell_offset = 16
            for item in defects['Сетка трещин'] or []:
                if check_in(item, (offset, offset + 100)):
                    row['defects'].append({
                        'type': 'Сетка трещин',
                        'short': "+",
                        "cell": defect_cell_offset + 1,
                        'alone': False,
                        'score': 1.9,
                        'length': item[1] + item[0],
                        'description': "сетка трещин на площади более 10 кв. м. "
                                       "при относительной площади занимаемой сеткой 60-30%",
                        'defect_code': '16',
                    })

            if not skip_koleynost:
                defect_cell_offset = 19
                for item in defects['Колейность'] or []:
                    if check_in(item, (offset, offset + 100)):
                        if item[2] == None:
                            cell_offset = 6
                        else:
                            cell_offset = [10, 20, 30, 40, 50, 70].index(item[2])

                        defect_item = {
                            10: {"description": "до 10мм", "code": '-', "score": 5.0},
                            20: {"description": "10-20мм", "code": '-', "score": 4.0},
                            30: {"description": "20-30мм", "code": '30', "score": 3.0},
                            40: {"description": "30-40мм", "code": '31', "score": 2.5},
                            50: {"description": "40-50мм", "code": '32', "score": 2.0},
                            70: {"description": "50-70мм", "code": '33', "score": 1.8},
                        }.get(item[2])

                        row['defects'].append({
                            'type': 'Колейность',
                            'short': "+",
                            'cell': defect_cell_offset + cell_offset,
                            'alone': item[2] is None,
                            'address': (item[0] + item[1]) / 2,
                            'length': item[1] + item[0],
                            'score': defect_item['score'],
                            'description': "колейность при средней глубине колеи {}".format(
                                defect_item['description']
                            ),
                            'defect_code': defect_item['code']
                        })

            defect_cell_offset = 32
            for item in defects['Выбоины'] or []:
                if check_in(item, (offset, offset + 100)):
                    if item[2] == None:
                        cell_offset = 0
                    else:
                        cell_offset = [20, 10, 4].index(item[2]) + 1

                    defect_item = {
                        None: {
                            "description": "одиночные выбоины на покрытиях, содержащих органическое вяжущее (расстояние между выбоинами более 20м)",
                            "code": '24',
                            "score": 4.5
                        },
                        20: {
                            "description": "одиночные выбоины на покрытиях, содержащих органическое вяжущее (расстояние между выбоинами 10-20м)",
                            "code": '25',
                            "score": 3.5
                        },
                        10: {
                            "description": "редкие выбоины на покрытиях, содержащих органическое вяжущее (расстояние между выбоинами 4-10м)",
                            "code": '26',
                            "score": 2.7
                        },
                        4: {
                            "description": "частые выбоины на покрытиях, содержащих органическое вяжущее (расстояние между выбоинами 1-4м)",
                            "code": '27',
                            "score": 2.2
                        },
                    }.get(item[2])

                    row['defects'].append({
                        'type': 'Выбоины',
                        'short': "+",
                        'cell': defect_cell_offset + cell_offset,
                        'alone': item[2] is None,
                        'address': (item[0] + item[1]) / 2,
                        'length': item[1] + item[0],
                        'score': defect_item['score'],
                        'description': defect_item['description'],
                        'defect_code': defect_item['code']
                    })

            defect_cell_offset = 33
            for item in defects['Карты(Заплаты)'] or []:
                if check_in(item, (offset, offset + 100)):
                    row['defects'].append({
                        'type': 'Карты(Заплаты)',
                        'short': "+",
                        "cell": defect_cell_offset,
                        'alone': False,
                        'score': 3.0,
                        'length': item[0] - item[1],
                        'description': "Карты заделанных выбоин, залитые трещины",
                        'defect_code': '28',
                    })

            if round_score:
                length = sum([i['length'] for i in row['defects'] if i['length']])
                score = sum([i['score'] * i['length'] for i in row['defects'] if i['length']])
                row['score'] = round(score / length, 1) if length else 0

            self.defects.append(row)

        return self.defects

    def get_barrier_data(self):
        if self.barrier_data:
            return self.barrier_data

        report = BarringReport(self.session)
        data = report(self.road.id, self.start, self.end)
        self.barrier_data = data

        return self.barrier_data

    def get_width_data(self):
        """
        Проверка кромки:
SELECT t.count, R.ID_Road, rtrim(R.Name)
FROM (
  SELECT
    ID_High,
    count(*) as count
  FROM Attribute
  WHERe ID_Type_Attr = 010108 -- and ID_High = 2-- ID_Attribute = 976541
  GROUP BY ID_High
) t
JOIN High h ON t.ID_High = h.ID_High
JOIN Way W on h.ID_Way = W.ID_Way
JOIN Road R on W.ID_Road = R.ID_Road
ORDER BY 1
        """
        attributes = Attribute.query_by_road(self.session, self.road.id).filter(
            Attribute.ID_Type_Attr == '010108'
        )
        attributes = list(attributes)

        # get periods by
        kromka_periods = RangeCustom(
            min=max(0, self.start),
            max=self.end,
            join_function=lambda x, y: (x or set()).union({y})
        )

        for idx, a in enumerate(attributes):
            points = sorted(a.points, key=lambda x: x.l)
            previous_point = points[0]
            for p in points[1:]:
                kromka_periods.add_subrange(previous_point.l, p.l, idx)
                previous_point = p

        rng = RangeCustom(
            min=max(0, self.start),
            max=self.end,
            join_function=lambda x, y: (x or 0) + (y or 0)
        )

        for start, end, idxs in kromka_periods.ranges:
            if idxs:
                idxs = list(idxs)
                if len(idxs) == 2:
                    left_points = attributes[idxs[0]].points
                    right_points = attributes[idxs[1]].points
                    left_points.sort(key=lambda x: x.l)
                    right_points.sort(key=lambda x: x.l)

                    if left_points[0].a >= right_points[0].a:
                        left_points, right_points = right_points, left_points

                    for points in (left_points, right_points):
                        previous_point = points[0]
                        for p in points[1:]:
                            value = abs((((previous_point.a + p.a) / 2) // 0.5) * 0.5)
                            rng.add_subrange(
                                max(start, previous_point.l),
                                min(end, p.l),
                                value,
                            )
                            previous_point = p

        # remove duplicates
        out_range = Range(
            min=max(0, self.start),
            max=self.end,
        )

        for r in rng.ranges:
            out_range.add_subrange(*r)

        return [r for r in out_range.ranges if r[2] is not None]

    def fill_table_trail_defects(self, table):
        report = TrailReport(self.session)
        data = report(self.road.id)
        for item in data:
            row = add_row(table, [
                item['type'],
                get_km(int((item['start'] + item['end']) / 2)),
                item['defect'],
            ])

    def fill_table_defects_by_odn(self, table):
        """Сводная ведомость наличия или отсутствия дефектов на участках автомобильных дорог"""
        defects = self.get_defects()

        table.rows[0].cells[0].text = self.road.Name

        for idx, data_row in enumerate(defects):
            row = table.add_row()
            cells = row.cells
            cells[1].text = data_row['address']
            for row_item in data_row['defects']:
                cells[row_item['cell']].text = row_item['short']

        if len(defects) > 3:
            cell = table.rows[3].cells[0]
            cell.merge(table.rows[2 + len(defects)].cells[0])
            cell.text = self.road.Name
            set_vertical_cell_direction(cell, 'tbRl')

    def fill_table_defects_by_odn_verbose(self, table):
        defects = self.get_defects()

        for idx, data_row in enumerate(defects):
            row = add_row(table, [
                self.road.Name,
                data_row['address'],
                ", ".join(list({i['description'] for i in data_row['defects'] if not i['alone']})),
                ", ".join(list({"{} ({})".format(
                    i['description'],
                    get_km(int(i['address']))
                ) for i in data_row['defects'] if i['alone']})),
                str(data_row['score']) if data_row['score'] else '-'
            ])

    def fill_table_barring_table(self, table):
        data = self.get_barrier_data()
        for row_item in data:
            row = add_row(table, [self.road.Name, row_item['Участок'], row_item['Тип'], row_item['condition']])

    def fill_bad_wells_table(self, table):
        report = BadWellsReport(self.session)
        data = report(self.road.id)
        for row_item in data:
            row = add_row(table, [row_item[0], get_km(row_item[1]), row_item[2]])

    def fill_totals(self, table):
        defects = self.get_defects()
        barriers = self.get_barrier_data()
        info = self.get_road_type_and_category()

        row = table.add_row()
        row.cells[0].text = self.road.Name
        row.cells[1].text = info['category']  # категория
        row.cells[2].text = str(round(float(self.end - max(0, self.start)) / 1000, 3))  # Протяженность, км

        score_cols = [0, 0, 0, 0, 0]
        potholes_count = 0
        other_alone_defects_count = 0
        for data_row in defects:
            score_cols[int(data_row['score']) - 1] += data_row['length']
            for defect in data_row['defects']:
                if defect['alone']:
                    if defect['type'] == 'Выбоины':
                        potholes_count += 1
                    else:
                        other_alone_defects_count += 1

        for idx, value in enumerate(score_cols):
            row.cells[3 + 4 - idx].text = str(round(value / 1000, 3))  # Протяженность участков покрытия, с оценкой состояния, км

        row.cells[8].text = str(potholes_count)  # Одиночные выбоины и проломы ДО, шт
        row.cells[9].text = str(other_alone_defects_count)  # Прочие точечные дефекты, шт

        barrier_good = 0
        barrier_bad = 0
        barrier_range = Range(0, self.end)

        for b in barriers:
            if b['is_barrier']:
                barrier_range.add_subrange(b['start'], b['end'], 1)
                if b['condition'] == 'хорошее':
                    barrier_good += b['length']
                else:
                    barrier_bad += b['length']

        barrier_without = sum([i[1] - i[0] for i in barrier_range.ranges if i[2] is None])

        # Протяженность участков ограждений, км
        row.cells[10].text = str(round(barrier_good / 1000, 3))
        row.cells[11].text = str(round(barrier_bad / 1000, 3))
        row.cells[12].text = str(round(barrier_without / 1000, 3))

    def fill_width_data(self, table):
        info = self.get_road_type_and_category()
        ranges = self.get_width_data()

        min_width = min([i[2] for i in ranges])
        max_width = max([i[2] for i in ranges])

        start = min([i[0] for i in ranges])
        end = max([i[1] for i in ranges])

        add_row(table, [
            self.road.Name,
            "{} — {}".format(get_km(int(start)), get_km(int(end))),
            "{} — {}".format(min_width, max_width),
            "—",
            "2",
            info['category'],
        ])

    def fill_not_normativ(self, table):
        smoothnes = self.get_smooth_data()

        for a in smoothnes:
            row = add_row(table, [
                "{} — {}".format(get_km(a['pos']), get_km(a['pos'] + a['delta'])),
                'не соответствует' if a['is_bad'] else 'соответствует',
                'Значение показателя продольной ровности '
                'покрытия по индексу IRI более требований, '
                'указанных в табл. 1 ГОСТ 33220-2015' if a['is_bad'] else '-',
            ])

    def fill_not_normativ_works(self, table):
        defects = self.get_defects()
        smoothnes = self.get_smooth_data()
        smoothnes = {i['pos']: i for i in smoothnes}

        previous_smooth = {}
        for idx, data_row in enumerate(defects):
            smooth = smoothnes.get(data_row['pos'])
            if smooth:
                previous_smooth = smooth
            else:
                smooth = previous_smooth

            if smooth and smooth['is_bad']:
                row = add_row(table, [
                    "{} — {}".format(get_km(data_row['pos']), get_km(data_row['pos'] + data_row['length'])),
                    'Капитальный ремонт участка улицы' if data_row['score'] < 3 else 'Ремонт покрытия проезжей части'
                ])

    def fill_total_paragpraph(self, paragraph1, paragraph2):
        smoothnes = self.get_smooth_data()

        if all([i['is_bad'] for i in smoothnes]):
            text = "Техническое и эксплуатационное состояния автомобильной дороги " \
                   "не соответствует существующему режиму движения"
            text2 = "Техническое и эксплуатационное состояния автомобильной дороги на всем протяжении " \
                    "не соответствует установленной категории автомобильной дороги"
        elif all([not i['is_bad'] for i in smoothnes]):
            text = "Техническое и эксплуатационное состояния автомобильной дороги " \
                   "соответствует существующему режиму движения"
            text2 = "Техническое и эксплуатационное состояния автомобильной дороги на всем протяжении " \
                    "соответствует установленной категории автомобильной дороги"
        else:
            text = "Техническое и эксплуатационное состояния автомобильной дороги " \
                   "на некоторых участках не соответствует существующему режиму движения."
            text2 = "Техническое и эксплуатационное состояния автомобильной дороги " \
                    "на некоторых участках не соответствует установленной категории автомобильной дороги"

        paragraph1.text = text
        paragraph2.text = text2

    def create_smooth_only(self):
        doc = docx.Document("templates/smooth.docx")
        self.progressed.emit(0, 1, "Заполняю таблицу по ровности")
        table = doc.tables[0]

        smoothnes = self.get_smooth_data()

        for a in smoothnes:
            row = add_row(table, [
                get_km(a['pos']),
                get_km(a['pos'] + a['delta']),
                str(a['forward']),
                str(a['backward']),
            ])

        self.progressed.emit(1, 1, "Готово")

        return doc

    def create_quality_only(self):
        doc = docx.Document("templates/quality.docx")
        self.progressed.emit(0, 1, "Заполняю таблицу по ровности")

        table = doc.tables[0]
        defects = self.get_defects()

        for idx, data_row in enumerate(defects):
            row = add_row(table, [
                get_km(data_row['pos']),
                get_km(data_row['pos'] + data_row['length']),
                str(data_row['score']) if data_row['score'] else str(5)
            ])

        self.progressed.emit(1, 1, "Готово")

        return doc

    def create_correspondence_only(self):
        doc = docx.Document("templates/correspondence.docx")
        self.progressed.emit(0, 1, "Заполняю таблицу соответствия")

        table = doc.tables[0]
        smoothnes = self.get_smooth_data()
        smoothnes = {i['pos']: i for i in smoothnes}
        defects = self.get_defects()

        previous_smooth = None

        for idx, data_row in enumerate(defects):
            smooth = smoothnes.get(data_row['pos'])
            if smooth:
                previous_smooth = smooth
            else:
                smooth = previous_smooth

            row = add_row(table, [
                get_km(data_row['pos']),
                get_km(data_row['pos'] + data_row['length']),
                'не соответствует' if smooth and smooth['is_bad'] else '-',
                '?',
                'не соответствует' if data_row['defects'] else '-'
            ])

        self.progressed.emit(1, 1, "Готово {}".format(self.road.Name))

        return doc

    def create_shelehov(self):
        doc = docx.Document("templates/sheleov.docx")

        table = doc.tables[0]

        ranges = self.get_width_data()
        min_width = min([i[2] for i in ranges])
        max_width = max([i[2] for i in ranges])
        start = min([i[0] for i in ranges])
        end = max([i[1] for i in ranges])
        add_row(table, [
            "{} — {}".format(get_km(int(start)), get_km(int(end))),
            "{} — {}".format(min_width, max_width),
        ])

        attributes = Attribute.query_by_road(self.session, self.road.id).filter(
            Attribute.ID_Type_Attr == '23120202',
        ).join(Params).with_entities(
            Attribute.L1,
            Attribute.L2,
            Attribute.ID_Type_Attr,
            Attribute.Image_Points,
            Attribute.Image_Counts,
            Params.value.label("value")
        ).order_by(Attribute.L1, Attribute.L2, Attribute.id)

        attributes = list(attributes)

        defects = self.get_defects_report()

        table = doc.tables[1]

        for p in defects.get('Сетка трещин', []) or []:
            add_row(table, [
                str(get_km(p[0])),
                str(get_km(p[1])),
                'Сетка трещин',
                str(round(p[2], 2)) + ' м2'
            ])

        delta = 50
        potholes = OrderedDict()
        for p in defects.get('Выбоины_raw', []) or []:
            potholes.setdefault(int(p[0] / delta), 0)
            potholes[int(p[0] / delta)] += p[2]

        for address, area in potholes.items():
            add_row(table, [
                str(get_km(address * delta)),
                str(get_km(address * delta + delta)),
                'Выбоины',
                str(round(area, 2)) + ' м2'
            ])

        cracks = OrderedDict()
        for p in defects.get('Трещины', []) or []:
            cracks.setdefault(int(p[0] / delta), 0)
            cracks[int(p[0] / delta)] += p[2]

        for address, area in cracks.items():
            add_row(table, [
                str(get_km(address * delta)),
                str(get_km(address * delta + delta)),
                'Трещины',
                str(round(area, 2)) + ' м'
            ])

        report = BadWellsReport(self.session)
        data = report(self.road.id)

        return doc

    def create(self):
        doc = docx.Document("templates/diagnostics.docx")

        count = 10

        self.fill_total_paragpraph(doc.paragraphs[3], doc.paragraphs[4])

        self.progressed.emit(0, count, "Заполняю не отвечающих нормативным требованиям таблицу")
        table = doc.tables[0]
        self.fill_not_normativ(table)

        self.progressed.emit(1, count, "Заполняю мероприятия по участкам не отвечающим нормативным требованиям таблицу")
        table = doc.tables[1]
        self.fill_not_normativ_works(table)

        self.progressed.emit(2, count, "Заполняю итоговую таблицу")
        table = doc.tables[2]
        self.fill_totals(table)

        self.progressed.emit(3, count, "Заполняю таблицу по дефектам")
        table = doc.tables[3]
        self.fill_table_defects_by_odn_verbose(table)

        self.progressed.emit(4, count, "Заполняю итоговую таблицу по дефектам")
        table = doc.tables[4]
        self.fill_table_defects_by_odn(table)

        self.progressed.emit(5, count, "Заполняю таблицу по ограждениям")
        table = doc.tables[5]
        self.fill_table_barring_table(table)

        self.progressed.emit(6, count, "Заполняю cводную ведомость категорий")
        table = doc.tables[6]
        self.fill_width_data(table)

        self.progressed.emit(7, count, "Заполняю таблицу по колодцам")
        table = doc.tables[7]
        self.fill_bad_wells_table(table)

        self.progressed.emit(8, count, "Заполняю таблицу по трамвайным путям")
        table = doc.tables[8]
        self.fill_table_trail_defects(table)

        self.progressed.emit(9, count, "Заполняю таблицу по ровности")
        table = doc.tables[9]
        self.fill_smooth_data(table)

        self.progressed.emit(count, count, "Готово")

        return doc

    def create_json(self):
        out = {
            "bad_wheels":[],
            'quality': [],
            'barrier': [],
            'trails': [],
        }

        report = BadWellsReport(self.session)
        data = report(self.road.id)
        for row_item in data:
            points = Attribute.get_points(row_item[3], row_item[4])
            out['bad_wheels'].append({
                "x": points[0].x,
                "y": points[0].y,
                "l": points[0].l,
                "type": row_item[0],
                "position": row_item[1],
                "defects": row_item[2],
            })

        # out['barrier'] = self.get_barrier_data()

        report = TrailReport(self.session)
        data = report(self.road.id)
        for item in data:
            out['trails'].append({
                'type': item['type'],
                'x': item['x'],
                'y': item['y'],
                'start': item['start'],
                'end': item['end'],
                'defect': item['defect'],
            })

        report = BarringReport(self.session)
        for b in report(self.road.id, self.start, self.end):
            out['barrier'].append({
                "range": b['Участок'],
                "position": b['Позиция'],
                "type": b['Тип'],
                "condition": b['condition'],
                "length": b['length'],
                'start': b['start'],
                'end': b['end'],
                'points': b['points'],
                'is_barrier': b['is_barrier'],
            })

        defects = self.get_defects()
        for idx, data_row in enumerate(defects):
            row = out['quality'].append({
                'pos': data_row['pos'],
                'length': data_row['length'],
                'defects': ", ".join(list({i['description'] for i in data_row['defects'] if not i['alone']})),
                'alonedefects': ''", ".join(list({"{} ({})".format(
                    i['description'],
                    get_km(int(i['address']))
                ) for i in data_row['defects'] if i['alone']})),
                'score': data_row['score']
            })

        return out


class DiagnosticsReportUlanUde2019(DiagnosticsReport):
    def set_table_header(self, table):
        table.rows[0].cells[0].paragraphs[0].text = self.road.Name
        table.rows[1].cells[0].paragraphs[0].text = "Дата обследования: 14.10.2019"

    def get_koleynost_data(self):
        out = []
        address_column = literal_column("L1 / {delta} * {delta}".format(delta=self.delta))

        def get_query(direction):
            return self.session.query(Attribute).filter(
                Attribute.ID_Type_Attr == '01020106',
                Attribute.id.in_(
                    Attribute.query_by_road(self.session, self.road.id)
                        .join(Params)
                        .filter(Params.id == '245')
                        .filter(Params.value == direction)
                        .with_entities(Attribute.id)
                )
            ).join(Params).filter(Params.id == '212')

        forward = get_query("Прямое").with_entities(
            address_column.label('pos'),
            literal_column("Avg(Cast(ValueParam as float))").label('forward'),
            literal_column('0').label('backward'),
        ).group_by(address_column)

        backward = get_query("Обратное").with_entities(
            address_column,
            literal_column("0"),
            literal_column("Avg(Cast(ValueParam as float))"),
        ).group_by(address_column)

        qs = forward.union(backward).subquery('t')

        attributes = self.session.query(
            qs.c.pos.label('pos'),
            case([(func.max(qs.c.forward) > 0, func.max(qs.c.forward))], else_=func.max(qs.c.backward)).label('forward'),
            case([(func.max(qs.c.backward) > 0, func.max(qs.c.backward)),], else_=func.max(qs.c.forward)).label('backward'),
        ).group_by(qs.c.pos)

        info = self.get_road_type_and_category()
        max_koleynost = {
            '1': 40,
            '1В': 40,
            '1B': 40,
            '2': 40,
            '3': 40,
            '4': 40,
            '5': 40,
        }.get(info['category'], 5.5)

        for a in attributes:
            if self.start <= a.pos <= self.end:
                out.append({
                    'backward': round(a.backward, 2),
                    'forward': round(a.forward, 2),
                    'delta': self.end - a.pos if a.pos + self.delta > self.end else self.delta,
                    'pos': a.pos,
                    'max_koleynost': max_koleynost,
                    'is_bad': max(a.backward, a.forward) > max_koleynost
                })

        return out

    def fill_table_defects_by_odn_verbose(self, table):
        self.set_table_header(table)

        defects = self.get_defects(round_score=False, skip_koleynost=True)
        self.road_length_defects_good = 0
        for defect_info in defects:

            cells = table.add_row().cells
            cells[0].text = str(defect_info['pos'] // 1000)
            cells[1].text = str(defect_info['pos'] % 1000)
            cells[2].text = str((defect_info['pos'] + defect_info['length']) // 1000)
            cells[3].text = str((defect_info['pos'] + defect_info['length']) % 1000)
            cells[4].text = str(defect_info['length'])

            if defect_info['defects']:
                min_defect_by_score = min(defect_info['defects'], key=lambda i: i['score'])
                is_good = min_defect_by_score['score'] >= 2.5

                cells[5].text = min_defect_by_score['defect_code']
                cells[6].text = min_defect_by_score['defect_code']
                cells[7].text = str(min_defect_by_score['score'])
                cells[8].text = str(min_defect_by_score['score'])
            else:
                is_good = True
                cells[5].text = '-'
                cells[6].text = '-'
                cells[7].text = '-'
                cells[8].text = '-'

            cells[9].text = "Соответствует" if is_good else "Не соответствует"

            key = (defect_info['pos'], defect_info['pos'] + defect_info['length'])
            self.length_good.setdefault(key, True)
            self.length_good[key] &= is_good

            if is_good:
                self.road_length_defects_good += defect_info['length']

    def fill_smooth_data(self, table):
        self.set_table_header(table)
        smooth_items = self.get_smooth_data()

        self.road_length_smooth_good = 0
        for a in smooth_items:
            cells = table.add_row().cells
            cells[0].text = get_km(a['pos'])
            cells[1].text = get_km(a['pos'] + a['delta'])
            cells[2].text = str(a['backward'])
            cells[3].text = str(a['forward'])
            cells[10].text = str(max(a['backward'], a['forward']))
            cells[11].text = "менее {}".format(a['max_iri'])
            cells[12].text = str(a['delta'])
            cells[13].text = "Не соответствует" if a['is_bad'] else "Соответствует"

            key = (a['pos'], a['pos'] + a['delta'])
            self.length_good.setdefault(key, True)
            self.length_good[key] &= not a['is_bad']

            if not a['is_bad']:
                self.road_length_smooth_good += a['delta']

    def create(self):
        doc_template = DocxTemplate("templates/ulan_ude_2019.docx")
        doc = doc_template.docx

        self.length_good = {}

        table = doc.tables[2]
        self.fill_smooth_data(table)

        # table = doc.tables[3]
        # self.fill_koleynost_data(table)

        table = doc.tables[3]
        self.fill_table_defects_by_odn_verbose(table)

        road_length_good = 0
        for key, value in self.length_good.items():
            if value:
                road_length_good += key[1] - key[0]

        doc_template.render({
            'road_name': self.road.Name,
            'road_length_smooth_good':  "{} км".format(self.road_length_smooth_good / 1000),
            'road_length_defects_good':  "{} км".format(self.road_length_defects_good / 1000),
            'road_length_good':  "{} км".format(road_length_good / 1000),
        })

        return doc

